"""DocxParser — .docx를 섹션 단위로 분리한다.

H1 1~4번(ISSUE, 요구사항, 사용자 관점, 개발자 관점)만 처리.
H1 5번째(사용자 및 운영자 매뉴얼) 이후는 스킵.
"""
from __future__ import annotations
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table

from src.parser.meta_extractor import extract_meta


# ── 데이터 클래스 ─────────────────────────────────────────────

@dataclass
class Section:
    """파싱된 섹션 하나."""
    section_id: str = ""          # "{DR번호}::{heading_path}"
    dr_number: str = ""
    title: str = ""               # 문서 제목
    heading_number: str = ""      # "3.1.1" 같은 번호
    heading_title: str = ""       # 헤딩 텍스트
    heading_path: str = ""        # "사용자 관점 > 업무규칙 정의 > 과제분석"
    level: int = 0                # 1~5
    section_type: str = ""        # 과제분석/구현방안/검증방안/DB참조/default
    content: str = ""             # 본문 텍스트 (테이블 마크다운 포함)
    content_length: int = 0
    tables_count: int = 0
    order: int = 0                # 문서 내 순서


@dataclass
class ParseResult:
    """파싱 결과."""
    meta: dict[str, Any] = field(default_factory=dict)
    sections: list[Section] = field(default_factory=list)


# ── 유틸 ──────────────────────────────────────────────────────

SECTION_TYPE_MAP = {
    "과제분석": "과제분석",
    "구현방안": "구현방안",
    "검증방안": "검증방안",
    "이슈사항": "이슈사항",
    "db object": "DB참조",
    "참조 레퍼런스": "DB참조",
    "기준정보": "DB참조",
}

ALLOWED_H1_COUNT = 4  # 처음 4개 H1만 처리


def _detect_section_type(heading_path: str) -> str:
    """heading_path에서 section_type을 추론."""
    lower = heading_path.lower()
    for keyword, stype in SECTION_TYPE_MAP.items():
        if keyword.lower() in lower:
            return stype
    return "default"


def _heading_level(style_name: str) -> int | None:
    """스타일명에서 헤딩 레벨 추출. 헤딩 아니면 None."""
    m = re.match(r"Heading\s*(\d)", style_name)
    return int(m.group(1)) if m else None


def _table_to_markdown(table: Table) -> str:
    """docx 테이블을 Markdown 테이블로 변환."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
    if len(rows) >= 2:
        # 헤더 구분선 삽입
        header_sep = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
        rows.insert(1, header_sep)
    return "\n".join(rows)


def _build_heading_number(level_counters: dict[int, int], level: int) -> str:
    """현재 레벨 카운터로 '3.1.1' 같은 번호 생성."""
    level_counters[level] = level_counters.get(level, 0) + 1
    # 하위 레벨 카운터 리셋
    for lv in list(level_counters.keys()):
        if lv > level:
            del level_counters[lv]
    parts = []
    for lv in range(1, level + 1):
        parts.append(str(level_counters.get(lv, 0)))
    return ".".join(parts)


# ── 메인 파서 ─────────────────────────────────────────────────

def parse_docx(file_path: str | Path) -> ParseResult:
    """docx 파일을 파싱하여 메타정보 + 섹션 리스트를 반환한다."""
    file_path = Path(file_path)
    doc = Document(str(file_path))

    # ── 1) 메타 테이블 추출 (Heading 전 테이블들) ──
    # body elements를 순회해서 첫 Heading 전의 테이블 수집
    meta_tables: list[Table] = []
    body = doc.element.body
    table_iter = iter(doc.tables)

    for element in body:
        tag = element.tag.split("}")[-1]
        if tag == "p":
            # 헤딩인지 확인
            pPr = element.find(qn("w:pPr"))
            if pPr is not None:
                pStyle = pPr.find(qn("w:pStyle"))
                if pStyle is not None:
                    val = pStyle.get(qn("w:val"), "")
                    if re.match(r"Heading\s*\d", val) or val.isdigit():
                        break
        elif tag == "tbl":
            try:
                meta_tables.append(next(table_iter))
            except StopIteration:
                break

    meta = extract_meta(meta_tables[:5], str(file_path))

    # ── 2) 헤딩 기반 섹션 분리 ──
    sections: list[Section] = []
    heading_stack: list[tuple[int, str]] = []  # (level, title)
    level_counters: dict[int, int] = {}
    h1_count = 0
    parsing_active = False
    stop_parsing = False

    current_content_parts: list[str] = []
    current_tables_count = 0
    current_heading_title = ""
    current_level = 0
    current_heading_number = ""
    current_order = 0

    def _flush_section():
        """현재 버퍼를 Section으로 저장."""
        nonlocal current_content_parts, current_tables_count
        if not current_heading_title:
            return
        content = "\n".join(current_content_parts).strip()
        if not content and current_tables_count == 0:
            # 빈 섹션은 스킵
            current_content_parts = []
            current_tables_count = 0
            return

        # heading_path 생성
        path_parts = [h[1] for h in heading_stack]
        heading_path = " > ".join(path_parts)

        section = Section(
            section_id=f"{meta['dr_number']}::{meta.get('target_year_month', '')}::{meta.get('part', '')}::{heading_path}",
            dr_number=meta["dr_number"],
            title=meta["title"],
            heading_number=current_heading_number,
            heading_title=current_heading_title,
            heading_path=heading_path,
            level=current_level,
            section_type=_detect_section_type(heading_path),
            content=content,
            content_length=len(content),
            tables_count=current_tables_count,
            order=current_order,
        )
        sections.append(section)
        current_content_parts = []
        current_tables_count = 0

    # body 재순회
    table_idx = 0
    all_tables = doc.tables

    for element in body:
        if stop_parsing:
            break

        tag = element.tag.split("}")[-1]

        if tag == "p":
            style_name = ""
            pPr = element.find(qn("w:pPr"))
            if pPr is not None:
                pStyle = pPr.find(qn("w:pStyle"))
                if pStyle is not None:
                    style_name = pStyle.get(qn("w:val"), "")

            # 헤딩 레벨 확인 — 숫자만 있는 스타일(예: "4")도 Heading으로 처리
            heading_lv = None
            if re.match(r"Heading\s*(\d)", style_name):
                heading_lv = int(re.search(r"(\d)", style_name).group(1))
            elif style_name.isdigit() and 1 <= int(style_name) <= 9:
                heading_lv = int(style_name)

            if heading_lv is not None:
                # 텍스트
                text = ""
                for child in element.iter():
                    if child.tag.endswith("}t"):
                        text += child.text or ""
                text = text.strip()

                if heading_lv == 1:
                    h1_count += 1
                    if h1_count <= ALLOWED_H1_COUNT:
                        parsing_active = True
                    else:
                        # 5번째 H1 → 파싱 종료
                        _flush_section()
                        stop_parsing = True
                        break

                if not parsing_active:
                    continue

                # 이전 섹션 저장
                _flush_section()

                # heading_number 생성
                current_heading_number = _build_heading_number(level_counters, heading_lv)
                current_heading_title = text
                current_level = heading_lv
                current_order += 1

                # 스택 업데이트
                while heading_stack and heading_stack[-1][0] >= heading_lv:
                    heading_stack.pop()
                heading_stack.append((heading_lv, text))

            else:
                # 일반 텍스트
                if not parsing_active:
                    continue
                # toc, Subtitle 등은 스킵
                if style_name.startswith("toc") or style_name == "Subtitle":
                    continue
                text = ""
                for child in element.iter():
                    if child.tag.endswith("}t"):
                        text += child.text or ""
                text = text.strip()
                if text:
                    current_content_parts.append(text)

        elif tag == "tbl":
            if table_idx < len(all_tables):
                tbl = all_tables[table_idx]
                table_idx += 1
                if parsing_active and current_heading_title:
                    md = _table_to_markdown(tbl)
                    current_content_parts.append(md)
                    current_tables_count += 1

    # 마지막 섹션 flush
    if not stop_parsing:
        _flush_section()

    return ParseResult(meta=meta, sections=sections)
